import re
import pandas as pd
from collections import Counter
from nltk.tokenize import word_tokenize
import random
import numpy as np
from scipy.stats import entropy
from nltk.corpus import stopwords
import os
# from sklearn.feature_extraction.text import TfidfVectorizer
import nltk
from nltk.corpus import stopwords
# import spacy
# from gensim import corpora, models
import json
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from langchain_google_genai import GoogleGenerativeAIEmbeddings


import docx
import os

def convert_docx_with_tables_to_text(docx_path):
    """
    Extracts text from a .docx file, correctly handling both paragraphs and tables.
    This is the corrected version that avoids the '.body' attribute error.

    Args:
        docx_path (str): The file path to the .docx file.

    Returns:
        str: The extracted and formatted text, or an empty string if an error occurs.
    """
    try:
        document = docx.Document(docx_path)
        full_text = []

        # First, extract text from all paragraphs
        for para in document.paragraphs:
            full_text.append(para.text)

        # Then, extract text from all tables
        for table in document.tables:
            # Add a marker to show where a table starts
            full_text.append("\n--- TABLE START ---")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                # Join cells with a separator for readability
                full_text.append("\t|\t".join(row_text))
            # Add a marker to show where a table ends
            full_text.append("--- TABLE END ---\n")

        return "\n".join(full_text)

    except Exception as e:
        print(f"Error reading DOCX file '{docx_path}': {e}")
        return ""